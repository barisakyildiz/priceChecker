import pandas as pd
import requests
from bs4 import BeautifulSoup
from price_parser import Price
import driver
import docx

class FileOp:
    def __init__(self):
        self.url_docx_name = "dataset.docx"
        self.url_file_name = "producturls.csv"
        self.prices_file_name = "prices.csv"
        self.SAVE_TO_CSV = True
        self.PRICES_CSV = "prices.csv"
        self.url = ""
    
    def readFromDoc(self):
        temp_total = []
        doc = docx.Document(self.url_docx_name)
        len_doc = len(doc.paragraphs)
        temp = []; flag = 0
        for i in range(len_doc):
            string = str(doc.paragraphs[i].text)
            for t in string.split():
                try:
                    temp.append(int(t))
                except ValueError:
                    pass
            for j in range(len(doc.paragraphs[i].text)):
                if (doc.paragraphs[i].text)[j] == '\n':
                    flag += 1
            if flag > 1:
                temp_total.append(temp)
                temp = []
                flag = 0
                slashn_loc = (doc.paragraphs[i].text).find('\n')
        print(temp_total)
            
    
    def getFromUrl(self):
        url_dataframe = pd.read_csv(self.url_file_name)
        return url_dataframe
    
    def processProd(self):
        updated = []
        dataframe = self.getFromUrl()
        for prod in dataframe.to_dict("records"):
            print("BEFORE ---> " + str(prod))
            webop = WebOp(prod["url"])
            html = webop.getHTML(prod["url"])
            prod["price"] = webop.returnPrice(html)
            prod["alert"] = prod["price"] < prod["alert_price"]
            updated.append(prod)
            print("AFTER ---> " + str(prod) + "\n\n")
            del webop
        return pd.DataFrame(updated)


def main():
    fileop = FileOp()
    fileop.readFromDoc()

if __name__ == '__main__':
    main()